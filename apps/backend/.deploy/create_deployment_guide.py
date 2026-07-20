from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "X3CRM_VPS_Deployment_and_Operations_Guide.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "18324A"
MUTED = "65758B"
INK = "1F2937"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
PALE_GOLD = "FFF6D9"
GOLD = "8A6500"
RED = "9B1C1C"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    if sum(widths_dxa) != 9360:
        raise ValueError(f"Table widths must total 9360 DXA: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "B7C4D4")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="Calibri", size=None, color=INK, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(table) -> None:
    if table.rows:
        repeat_header(table.rows[0])


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_run_font(run, size=8.5, color=MUTED)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])


def add_numbering(doc: Document, fmt: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if fmt == "bullet" else "%1.")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "271")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])


def add_list_item(doc, text: str, num_id: int, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    apply_numbering(p, num_id)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_body(doc, text: str, bold_prefix: str | None = None, italic=False):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic)
    return p


def add_heading(doc, text: str, level: int):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_code(doc, text: str):
    for index, line in enumerate(text.strip("\n").splitlines()):
        p = doc.add_paragraph(style="Code Block")
        p.paragraph_format.space_before = Pt(4 if index == 0 else 0)
        p.paragraph_format.space_after = Pt(4 if index == len(text.strip("\n").splitlines()) - 1 else 0)
        p.paragraph_format.keep_together = True
        r = p.add_run(line or " ")
        set_run_font(r, name="Consolas", size=8.2, color="203040")


def add_callout(doc, label: str, text: str, fill=LIGHT_GRAY, color=INK):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "14")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    r = p.add_run(f"{label}: ")
    set_run_font(r, size=10.2, bold=True, color=color)
    r = p.add_run(text)
    set_run_font(r, size=10.2, color=color)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    header = table.rows[0]
    for index, text in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        set_run_font(r, size=9.4, color=NAVY, bold=True)
    repeat_header(header)
    for row_values in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row_values):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run(str(text))
            set_run_font(r, size=9.2)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(29)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.line_spacing = 1.0

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(14)
    subtitle.font.italic = False
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(10)

    heading_tokens = {
        1: (16, BLUE, 18, 10),
        2: (13, BLUE, 14, 7),
        3: (12, DARK_BLUE, 10, 5),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    if "Code Block" not in [style.name for style in styles]:
        code = styles.add_style("Code Block", 1)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(8.2)
    code.paragraph_format.left_indent = Inches(0.12)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.line_spacing = 1.05
    p_pr = code._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EEF2F6")
    p_pr.append(shd)


def configure_sections(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("X3 CRM  |  VPS Deployment & Operations")
    set_run_font(r, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    add_page_number(footer.paragraphs[0])


def build_document() -> Path:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    bullet_id = add_numbering(doc, "bullet")

    # Editorial cover pattern for a technical manual; intentionally no header border.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(88)
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RUNBOOK  |  VPS / DOCKER")
    set_run_font(r, size=10.5, bold=True, color=GOLD)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("X3 CRM").font.name = "Calibri"
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Hướng dẫn triển khai, vận hành, cập nhật và khôi phục")
    set_run_font(r, size=14, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(32)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("VPS: 45.252.251.120")
    set_run_font(r, size=12, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Trạng thái ghi nhận: 18/07/2026 • Môi trường chạy thử")
    set_run_font(r, size=10.5, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    r = p.add_run("Không chứa mật khẩu SSH, APP_KEY hoặc mật khẩu PostgreSQL.")
    set_run_font(r, size=9.5, color=RED, bold=True)
    doc.add_page_break()

    add_heading(doc, "1. Mục đích và phạm vi", 1)
    add_body(doc, "Tài liệu này ghi lại cấu hình đã triển khai lên VPS, quy trình đưa phiên bản mới lên máy chủ, cách sao lưu/khôi phục và các bước xử lý sự cố thường gặp. Đây là tài liệu vận hành cho môi trường chạy thử bằng địa chỉ IP; chưa phải cấu hình production hoàn chỉnh.")
    add_callout(doc, "Điểm truy cập", "Ứng dụng hiện chạy tại http://45.252.251.120. Chỉ cổng 80 (HTTP) và 22 (SSH) được công khai; PostgreSQL không mở cổng ra Internet.", fill=PALE_GOLD, color=GOLD)

    add_heading(doc, "2. Trạng thái hệ thống hiện tại", 1)
    add_table(doc, ["Hạng mục", "Giá trị đã triển khai"], [
        ["Máy chủ", "Ubuntu 24.04.3 LTS; hostname crm.x3sales.vn"],
        ["Tài nguyên", "1 vCPU; khoảng 709 MiB RAM; bổ sung 1 GiB swap"],
        ["Docker", "Docker Engine 29.1.3; Docker Compose 2.37.1"],
        ["Thư mục chạy", "/opt/x3crm"],
        ["URL", "http://45.252.251.120"],
        ["Reverse proxy", "Nginx 1.27 Alpine, public port 80"],
        ["Backend", "Laravel 11 / PHP 8.4 Alpine; internal port 4000"],
        ["Frontend", "Next.js 14 standalone / Node 20 Alpine; internal port 3000"],
        ["Database", "PostgreSQL 17 Alpine; internal port 5432; named volume db_data"],
        ["Uploads", "Named volume x3crm_uploads_data; Nginx phục vụ tại /uploads/"],
    ], [2700, 6660])

    add_heading(doc, "2.1 Luồng request", 2)
    add_code(doc, """
Browser -> 45.252.251.120:80 -> Nginx
  /                  -> Frontend :3000
  /api/, /sanctum/   -> Backend  :4000 -> PostgreSQL :5432
  /uploads/          -> volume x3crm_uploads_data (read-only tại Nginx)
""")
    add_body(doc, "Các container hiện tại: x3crm-nginx-1, x3crm-frontend-1, x3crm-backend-1 và x3crm-db-1. Tất cả được điều phối bởi Docker Compose project `x3crm`.")

    add_heading(doc, "2.2 Dữ liệu sau lần reset ngày 18/07/2026", 2)
    add_table(doc, ["Nhóm", "Kết quả sau reset"], [
        ["Tài khoản", "users = 5; roles = 5; permissions = 10; role_permissions = 10"],
        ["Phòng ban", "departments = 0 (giữ nguyên trạng thái trước reset)"],
        ["Dịch vụ", "services = 54; service_packages = 0 (giữ nguyên)"],
        ["Schema", "migrations = 64"],
        ["Dữ liệu nghiệp vụ", "Tất cả bảng ngoài danh sách giữ lại = 0"],
        ["Hình ảnh/uploads", "0 file; URL ảnh cũ trả HTTP 404"],
    ], [2700, 6660])
    add_callout(doc, "Lưu ý", "Bảng sessions đã được xóa trong lần reset. Người dùng phải đăng nhập lại, nhưng tài khoản và mật khẩu không bị thay đổi.")

    add_heading(doc, "3. Bộ file triển khai", 1)
    add_table(doc, ["File", "Vai trò"], [
        [".deploy/production/compose.yml", "Khai báo 4 dịch vụ, volumes, healthcheck và giới hạn log"],
        ["backend.Dockerfile", "Build Laravel/PHP 8.4, pdo_pgsql, OPcache và Composer no-dev"],
        ["frontend.Dockerfile", "Build Next.js standalone bằng Node 20"],
        ["nginx.conf", "Định tuyến frontend, API, Sanctum và uploads"],
        ["opcache.ini", "Cấu hình OPcache cho backend"],
        ["next.config.production.js", "Bật output standalone và đưa biến public vào build"],
        ["env.template", "Mẫu biến môi trường; không chứa secret thật"],
        ["reset-keep-accounts-services.sql", "Reset dữ liệu có kiểm soát, giữ tài khoản/dịch vụ"],
        [".dockerignore", "Loại trừ secret, vendor, node_modules, log và artifact khỏi build context"],
    ], [3300, 6060])
    add_body(doc, "Trên VPS, các file cấu hình được đặt trong `/opt/x3crm`. File `.env` thật chỉ nằm trên VPS, quyền nên là `600`. Không đưa `.env`, APP_KEY, mật khẩu database hoặc mật khẩu SSH vào Git và tài liệu này.")

    add_heading(doc, "4. Quy trình triển khai ban đầu đã thực hiện", 1)
    add_heading(doc, "4.1 Chuẩn bị VPS", 2)
    for item in [
        "Đăng nhập SSH bằng tài khoản root và xác minh đúng host/IP trước khi thay đổi.",
        "Cập nhật package index, cài Docker Engine và Docker Compose plugin từ repository chính thức của Docker.",
        "Tạo 1 GiB swap để giảm nguy cơ hết RAM khi container khởi động hoặc migrate.",
        "Tạo thư mục `/opt/x3crm` và volume uploads dùng chung.",
    ]:
        add_list_item(doc, item, bullet_id)
    add_code(doc, """
fallocate -l 1G /swapfile
chmod 600 /swapfile
mkswap /swapfile
swapon /swapfile
echo '/swapfile none swap sw 0 0' >> /etc/fstab

mkdir -p /opt/x3crm
docker volume create x3crm_uploads_data
""")

    add_heading(doc, "4.2 Build image trên máy phát triển", 2)
    add_body(doc, "Image được build cho `linux/amd64` trên máy local. Backend dùng source tại `apps/backend`; frontend dùng source tại `apps/frontend` cùng bản `next.config.production.js` dành cho production.")
    add_code(doc, """
# Từ apps/backend
docker build --platform linux/amd64 \\
  -f .deploy/production/backend.Dockerfile \\
  -t x3crm-backend:deploy .

# Từ build context frontend đã có next.config.production.js
docker build --platform linux/amd64 \\
  --build-arg NEXT_PUBLIC_API_URL=http://45.252.251.120/api \\
  --build-arg NEXT_PUBLIC_MEDIA_URL=http://45.252.251.120 \\
  -f frontend.Dockerfile -t x3crm-frontend:deploy .
""")
    add_callout(doc, "Quan trọng", "Các biến `NEXT_PUBLIC_*` được đóng vào frontend tại thời điểm build. Khi đổi IP hoặc trỏ domain, phải build lại image frontend; chỉ sửa `.env` trên VPS là chưa đủ.", fill=PALE_GOLD, color=GOLD)

    add_heading(doc, "4.3 Đóng gói và chuyển artifact", 2)
    add_code(doc, """
docker save -o x3crm-backend.tar x3crm-backend:deploy
docker save -o x3crm-frontend.tar x3crm-frontend:deploy

scp x3crm-backend.tar root@45.252.251.120:/opt/x3crm/
scp x3crm-frontend.tar root@45.252.251.120:/opt/x3crm/
scp .deploy/production/compose.yml root@45.252.251.120:/opt/x3crm/
scp .deploy/production/nginx.conf root@45.252.251.120:/opt/x3crm/
""")
    add_body(doc, "Trong lần triển khai đầu, database local được xuất bằng `pg_dump` định dạng custom và uploads được nén thành tar.gz, sau đó tải lên VPS. Các image được `docker load`, database/uploads được restore vào volumes trước khi khởi động toàn bộ stack.")

    add_heading(doc, "4.4 Tạo biến môi trường", 2)
    add_body(doc, "Copy `env.template` thành `/opt/x3crm/.env`, thay placeholder bằng APP_KEY và DB_PASSWORD ngẫu nhiên, sau đó đặt quyền file. Không dùng lại mật khẩu SSH làm mật khẩu database.")
    add_code(doc, """
cd /opt/x3crm
cp env.template .env
# Thay __APP_KEY__ và __DB_PASSWORD__ bằng giá trị bí mật riêng
chmod 600 .env
""")

    add_heading(doc, "4.5 Load image, restore và khởi động", 2)
    add_code(doc, """
cd /opt/x3crm
docker load -i x3crm-backend.tar
docker load -i x3crm-frontend.tar
docker compose up -d db

# Restore database custom dump
docker compose exec -T db pg_restore \\
  -U x3crm -d x3crm --no-owner --no-privileges < x3crm.dump

# Restore uploads vào volume dùng chung
docker run --rm \\
  -v x3crm_uploads_data:/target \\
  -v /opt/x3crm:/backup alpine:3.22 \\
  sh -c 'tar -xzf /backup/uploads.tar.gz -C /target'

docker compose up -d
docker compose ps
""")
    add_body(doc, "Backend entry command tự chạy `php artisan migrate --force`, `php artisan optimize`, rồi khởi động PHP server tại cổng 4000. Nginx chỉ mở cổng 80; database không publish cổng 5432.")

    add_heading(doc, "5. Kiểm tra sau deploy", 1)
    add_table(doc, ["Kiểm tra", "Kết quả mong đợi"], [
        ["docker compose ps", "4 container Up; database Healthy"],
        ["GET /", "HTTP 200, giao diện tải bình thường"],
        ["POST /api/auth/login", "Đăng nhập thành công với tài khoản hợp lệ"],
        ["GET /api/services chưa đăng nhập", "HTTP 401 với Accept: application/json"],
        ["Tải/tạo ảnh", "File xuất hiện dưới /uploads/ và trả HTTP 200"],
        ["docker compose logs", "Không có production.ERROR mới hoặc vòng restart"],
    ], [3300, 6060])
    add_code(doc, """
cd /opt/x3crm
docker compose ps
docker compose logs --since=10m --tail=200
curl -I http://127.0.0.1/
curl -H 'Accept: application/json' -i http://127.0.0.1/api/services
""")

    add_heading(doc, "6. Vận hành hằng ngày", 1)
    add_table(doc, ["Nhu cầu", "Lệnh"], [
        ["Xem trạng thái", "cd /opt/x3crm && docker compose ps"],
        ["Theo dõi log", "docker compose logs -f --tail=200"],
        ["Log một dịch vụ", "docker compose logs -f --tail=200 backend"],
        ["Khởi động stack", "docker compose up -d"],
        ["Dừng stack", "docker compose down"],
        ["Restart backend", "docker compose restart backend"],
        ["Chạy migrate", "docker compose exec -T backend php artisan migrate --force"],
        ["Kiểm tra tài nguyên", "docker stats --no-stream; df -h; free -h"],
    ], [3000, 6360])
    add_callout(doc, "Không dùng", "Không chạy `docker compose down -v` trong vận hành bình thường vì tùy chọn `-v` có thể xóa volume database. Không xóa `/opt/x3crm/.env` nếu chưa có bản sao secret.", fill=PALE_GOLD, color=RED)

    add_heading(doc, "7. Cập nhật phiên bản mới", 1)
    for item in [
        "Chạy test/build trên local và tạo backup database/uploads trước khi cập nhật.",
        "Build lại image backend khi PHP/Laravel thay đổi; build lại frontend khi source hoặc IP/domain public thay đổi.",
        "Gắn tag có version (ví dụ `x3crm-backend:2026-07-18-01`) thay vì chỉ dùng `deploy` nếu cần rollback nhanh.",
        "Chuyển image lên VPS, `docker load`, sửa tag trong compose.yml và chạy `docker compose up -d`.",
        "Kiểm tra migrate, login, chức năng chính, log và tài nguyên; chỉ xóa image cũ khi bản mới ổn định.",
    ]:
        add_list_item(doc, item, bullet_id)
    add_code(doc, """
cd /opt/x3crm
docker load -i x3crm-backend-NEW.tar
docker load -i x3crm-frontend-NEW.tar
docker compose up -d
docker compose ps
docker compose logs --since=5m --tail=200
""")
    add_body(doc, "Hiện source code không được checkout trực tiếp trên VPS; VPS chạy từ Docker images và các file cấu hình. Vì vậy quy trình thay đổi chuẩn là build trên máy phát triển, chuyển image và redeploy. Không chỉnh source thủ công trong container vì thay đổi sẽ mất khi container được tạo lại.")

    add_heading(doc, "8. Sao lưu", 1)
    add_heading(doc, "8.1 Sao lưu database", 2)
    add_code(doc, """
cd /opt/x3crm
ts=$(date +%Y%m%d-%H%M%S)
mkdir -p backups
docker compose exec -T db pg_dump \\
  -U x3crm -d x3crm -Fc --no-owner --no-privileges \\
  > "backups/x3crm-db-${ts}.dump"
chmod 600 "backups/x3crm-db-${ts}.dump"
sha256sum "backups/x3crm-db-${ts}.dump"
""")
    add_heading(doc, "8.2 Sao lưu uploads", 2)
    add_code(doc, """
cd /opt/x3crm
ts=$(date +%Y%m%d-%H%M%S)
docker run --rm \\
  -v x3crm_uploads_data:/source:ro \\
  -v /opt/x3crm/backups:/backup alpine:3.22 \\
  sh -c "tar -czf /backup/x3crm-uploads-${ts}.tar.gz -C /source ."
chmod 600 "backups/x3crm-uploads-${ts}.tar.gz"
""")
    add_body(doc, "Nên tải backup sang một nơi khác VPS (máy local hoặc object storage). Backup nằm cùng ổ đĩa với ứng dụng không bảo vệ được khi VPS hoặc ổ đĩa hỏng.")

    add_heading(doc, "8.3 Backup trước lần reset 18/07/2026", 2)
    add_table(doc, ["File trên VPS", "SHA-256"], [
        ["/opt/x3crm/backups/pre-reset-db-20260718-120112.dump", "784f4a59eed078e735eefecfac2bf25238b310134efe775489cc6f855917e6d7"],
        ["/opt/x3crm/backups/pre-reset-uploads-20260718-120112.tar.gz", "899f651de2e0004cfe2824e96cffb9b3fca1644d66848ed7842db0051b1a19ea"],
    ], [4100, 5260])

    add_heading(doc, "9. Reset dữ liệu, giữ tài khoản và dịch vụ", 1)
    add_body(doc, "Lần reset này đã sử dụng script `/opt/x3crm/reset-keep-accounts-services.sql`. Script thực hiện một transaction và TRUNCATE mọi bảng public ngoài danh sách giữ lại, đồng thời RESTART IDENTITY và CASCADE cho các bảng nghiệp vụ liên quan.")
    add_table(doc, ["Nhóm giữ lại", "Bảng"], [
        ["Schema", "migrations"],
        ["Tài khoản/phụ thuộc", "users, roles, permissions, role_permissions, departments"],
        ["Dịch vụ", "services, service_packages"],
    ], [2700, 6660])
    add_code(doc, """
cd /opt/x3crm

# 1) Bắt buộc backup trước
# 2) Chạy reset database
docker compose exec -T db psql \\
  -U x3crm -d x3crm \\
  < reset-keep-accounts-services.sql

# 3) Xóa toàn bộ file uploads trong volume
docker run --rm -v x3crm_uploads_data:/target alpine:3.22 \\
  sh -c 'find /target -mindepth 1 -delete'

# 4) Dọn config/route/view và restart backend
docker compose exec -T backend php artisan config:clear
docker compose exec -T backend php artisan route:clear
docker compose exec -T backend php artisan view:clear
docker compose restart backend
""")
    add_callout(doc, "Cảnh báo", "Không chạy lại reset nếu chưa xác nhận đúng VPS, đúng database và có backup đọc được. Lệnh này xóa dữ liệu nghiệp vụ thật và không thể undo nếu không restore backup.", fill=PALE_GOLD, color=RED)
    add_body(doc, "Không dùng `php artisan optimize:clear` ở cấu hình hiện tại vì Laravel đang chọn database cache nhưng schema không có bảng `cache`; lệnh tổng hợp sẽ báo `relation cache does not exist`. Dùng ba lệnh config:clear, route:clear và view:clear như trên.")

    add_heading(doc, "10. Khôi phục từ backup", 1)
    add_heading(doc, "10.1 Khôi phục database đầy đủ", 2)
    add_callout(doc, "Downtime", "Quy trình này dừng backend trong thời gian restore. Xác nhận tên file backup và checksum trước khi chạy.", fill=PALE_GOLD, color=GOLD)
    add_code(doc, """
cd /opt/x3crm
docker compose stop backend

docker compose exec -T db dropdb -U x3crm --if-exists --force x3crm
docker compose exec -T db createdb -U x3crm x3crm
docker compose exec -T db pg_restore \\
  -U x3crm -d x3crm --no-owner --no-privileges \\
  < backups/pre-reset-db-20260718-120112.dump

docker compose start backend
docker compose logs --since=5m --tail=200 backend
""")
    add_heading(doc, "10.2 Khôi phục uploads", 2)
    add_code(doc, """
cd /opt/x3crm
docker run --rm \\
  -v x3crm_uploads_data:/target \\
  -v /opt/x3crm/backups:/backup alpine:3.22 \\
  sh -c 'find /target -mindepth 1 -delete && \
         tar -xzf /backup/pre-reset-uploads-20260718-120112.tar.gz \
         -C /target'
""")
    add_body(doc, "Sau restore, kiểm tra số lượng bảng, login, trang danh sách, URL ảnh, log backend và trạng thái 4 container. Nếu APP_KEY đã đổi so với lúc backup, session/cookie cũ sẽ không còn hợp lệ; tài khoản/mật khẩu trong database vẫn giữ nguyên.")

    add_heading(doc, "11. Xử lý sự cố nhanh", 1)
    add_table(doc, ["Triệu chứng", "Kiểm tra / hướng xử lý"], [
        ["Trang không mở", "docker compose ps; kiểm tra Nginx; curl http://127.0.0.1/; kiểm tra firewall cổng 80"],
        ["API 502", "Backend đang restart/chưa listen; xem logs backend, chờ migrate/optimize hoàn tất"],
        ["API 500 khi chưa đăng nhập", "Gửi header Accept: application/json. Nếu thiếu, middleware có thể tìm route `login` không tồn tại"],
        ["Đăng nhập lỗi", "Kiểm tra users/roles, DB healthy, SANCTUM_STATEFUL_DOMAINS, cookie và thời gian hệ thống"],
        ["Ảnh 404", "Kiểm tra file trong x3crm_uploads_data, đường dẫn lưu DB và alias /uploads/ của Nginx"],
        ["DB connection refused", "docker compose ps db; healthcheck; DB_HOST phải là `db`, không phải localhost"],
        ["Container restart liên tục", "docker compose logs SERVICE; kiểm tra RAM/swap, disk, biến .env và migrate"],
        ["Ổ đĩa đầy", "df -h; docker system df; kiểm tra backup/log/image cũ trước khi dọn"],
        ["Đổi domain nhưng frontend gọi IP cũ", "Build lại frontend với NEXT_PUBLIC_API_URL và NEXT_PUBLIC_MEDIA_URL mới"],
    ], [3000, 6360])

    add_heading(doc, "12. Việc cần làm trước production", 1)
    for item in [
        "Trỏ domain và bật HTTPS; sau đó cập nhật APP_URL, FRONTEND_URL(S), SANCTUM_STATEFUL_DOMAINS và SESSION_SECURE_COOKIE=true.",
        "Đổi mật khẩu root đã chia sẻ, tạo user quản trị riêng, dùng SSH key và tắt password login/root login khi đã xác nhận key hoạt động.",
        "Thay PHP built-in server (`php -S`) bằng PHP-FPM + Nginx hoặc FrankenPHP. Cấu hình hiện tại phù hợp chạy thử, không phải application server production.",
        "Thiết lập firewall chỉ mở 22/80/443; giới hạn SSH theo IP nếu có thể.",
        "Tự động backup database/uploads ra nơi khác VPS và kiểm thử restore định kỳ.",
        "Bổ sung healthcheck cho backend/frontend, giám sát uptime, cảnh báo dung lượng/RAM và log tập trung.",
        "Dùng image tag theo phiên bản và lưu ít nhất một phiên bản ổn định để rollback.",
        "Tăng RAM/CPU nếu số người dùng hoặc tác vụ đồng thời tăng; VPS hiện có tài nguyên rất thấp.",
    ]:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "13. Checklist trước và sau mỗi lần thay đổi", 1)
    add_table(doc, ["Giai đoạn", "Checklist"], [
        ["Trước thay đổi", "Xác nhận đúng VPS; backup DB/uploads; ghi tag image cũ; kiểm tra disk/RAM"],
        ["Triển khai", "Load image; cập nhật compose/.env có kiểm soát; up -d; theo dõi migrate/log"],
        ["Xác minh", "4 container Up; frontend 200; login; API chính; upload; không có error mới"],
        ["Sau triển khai", "Ghi phiên bản/thời gian; giữ backup và image rollback; không xóa secret"],
    ], [2500, 6860])

    add_heading(doc, "14. Nhật ký thay đổi tài liệu", 1)
    add_table(doc, ["Ngày", "Thay đổi"], [
        ["18/07/2026", "Ghi nhận deployment Docker đầu tiên lên 45.252.251.120; bổ sung quy trình vận hành, reset và rollback."],
        ["18/07/2026", "Backup và reset: giữ tài khoản/phân quyền/dịch vụ; xóa dữ liệu nghiệp vụ, sessions và toàn bộ uploads."],
    ], [1800, 7560])

    add_callout(doc, "Kết luận", "Hệ thống hiện hoạt động ổn định ở chế độ chạy thử qua IP. Backup trước reset đã được giữ trên VPS; tài khoản và dịch vụ còn nguyên, dữ liệu nghiệp vụ và uploads đã được làm sạch.", fill=LIGHT_BLUE, color=NAVY)

    doc.core_properties.title = "X3 CRM - VPS Deployment and Operations Guide"
    doc.core_properties.subject = "Runbook triển khai và vận hành X3 CRM trên Docker"
    doc.core_properties.author = "X3 Sales"
    doc.core_properties.keywords = "X3 CRM, Docker, VPS, Laravel, Next.js, PostgreSQL, runbook"
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
